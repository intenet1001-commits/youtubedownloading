#!/usr/bin/env python3
"""
테스트용 DOCX 파일 생성 스크립트
표, 리스트, 서식이 포함된 문서를 생성하여 변환 테스트에 사용
"""

from docx import Document
from docx.shared import Inches

def create_test_document():
    # 새 문서 생성
    doc = Document()
    
    # 제목 추가
    title = doc.add_heading('문서 변환 테스트 문서', level=1)
    
    # 서론 문단
    intro = doc.add_paragraph('이 문서는 PDF, PPT, MD, 워드 간 변환 기능을 테스트하기 위해 작성되었습니다. ')
    intro.add_run('다양한 서식').bold = True
    intro.add_run('과 ')
    intro.add_run('구조적 요소').italic = True
    intro.add_run('들이 포함되어 있습니다.')
    
    # 부제목
    doc.add_heading('주요 기능', level=2)
    
    # 번호 있는 리스트
    doc.add_paragraph('1. PDF ↔ DOCX, MD, PPTX 변환')
    doc.add_paragraph('2. DOCX ↔ PDF, MD, PPTX 변환')  
    doc.add_paragraph('3. PPTX ↔ PDF, DOCX, MD 변환')
    doc.add_paragraph('4. MD ↔ PDF, DOCX, PPTX 변환')
    
    # 부제목
    doc.add_heading('변환 성능 비교표', level=2)
    
    # 표 추가
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 헤더 행
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '입력 형식'
    hdr_cells[1].text = '출력 형식'
    hdr_cells[2].text = '처리 시간'
    hdr_cells[3].text = '품질'
    
    # 데이터 행들
    data = [
        ('PDF', 'MD', '빠름', '높음'),
        ('DOCX', 'MD', '매우 빠름', '매우 높음'),
        ('PPTX', 'MD', '빠름', '높음'),
        ('MD', 'PDF', '보통', '높음')
    ]
    
    for input_fmt, output_fmt, speed, quality in data:
        row_cells = table.add_row().cells
        row_cells[0].text = input_fmt
        row_cells[1].text = output_fmt
        row_cells[2].text = speed
        row_cells[3].text = quality
    
    # 부제목
    doc.add_heading('특징 및 장점', level=2)
    
    # 불릿 포인트 리스트
    doc.add_paragraph('• 직관적인 사용자 인터페이스', style='List Bullet')
    doc.add_paragraph('• 실시간 변환 진행률 표시', style='List Bullet')
    doc.add_paragraph('• 다양한 문서 형식 지원', style='List Bullet')
    doc.add_paragraph('• 표와 서식 보존', style='List Bullet')
    
    # 결론
    doc.add_heading('결론', level=2)
    conclusion = doc.add_paragraph('이 ')
    conclusion.add_run('문서 변환 도구').bold = True
    conclusion.add_run('는 다양한 문서 형식 간의 ')
    conclusion.add_run('효율적인 변환').italic = True
    conclusion.add_run('을 제공합니다. 표, 리스트, 서식 등이 ')
    conclusion.add_run('정확히 보존').bold = True
    conclusion.add_run('되어 변환됩니다.')
    
    # 문서 저장
    output_path = '/Users/choichunsung/Documents/GitHub/myproduct_v4/youtubedownloading/test_document_with_table.docx'
    doc.save(output_path)
    print(f"테스트 문서가 생성되었습니다: {output_path}")

if __name__ == "__main__":
    create_test_document()